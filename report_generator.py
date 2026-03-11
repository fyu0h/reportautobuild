"""
Word 报告生成模块
================
使用 python-docx 在服务端生成符合格式规范的 Word 文档。
严格遵循《皇岗边检站国际移民一周资讯》的格式规范。
"""

import io
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# 中文数字
CN_NUMS = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']

# 行距 28pt = 固定值
LINE_SPACING_PT = 28
# 首行缩进约 2 字符（16pt 字体 × 2 = 32pt ≈ 1.13cm）
FIRST_LINE_INDENT = Cm(1.13)


def set_run_font(run, font_name_east, font_name_ascii="Times New Roman",
                  size_pt=16, bold=False, color_hex="000000"):
    """设置 run 的字体属性"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name_ascii
    # 设置东亚字体
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_east)
    if color_hex:
        run.font.color.rgb = None  # reset
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(color_hex)


def set_paragraph_spacing(paragraph, line_spacing_pt=LINE_SPACING_PT):
    """设置段落固定行距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = Pt(line_spacing_pt)
    pf.line_spacing_rule = 2  # EXACT (固定值)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def generate_word_report(report_data: dict) -> io.BytesIO:
    """
    根据结构化报告数据生成 Word 文档。

    report_data 格式:
    {
        "title": "皇岗边检站国际移民一周资讯",
        "dateStart": "3月3日",
        "dateEnd": "3月9日",
        "summaries": ["摘要1；", "摘要2；"],
        "sections": [
            {
                "title": "涉我重要移民动态",
                "newsItems": [
                    {"headline": "新闻标题。", "body": "据...报道...（皇岗边检站）"}
                ]
            }
        ]
    }

    返回 BytesIO 对象（可直接发送给客户端下载）。
    """
    doc = Document()

    # ===== 页面设置 =====
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.50)
    section.footer_distance = Cm(1.75)

    title = report_data.get("title", "皇岗边检站国际移民一周资讯")
    date_start = report_data.get("dateStart", "X月X日")
    date_end = report_data.get("dateEnd", "X月X日")
    summaries = report_data.get("summaries", [])
    sections = report_data.get("sections", [])

    # ===== 大标题 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p)
    run = p.add_run(title)
    set_run_font(run, "方正小标宋简体", "Times New Roman", 22, bold=False)

    # ===== 副标题（日期） =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p)
    run = p.add_run(f"（{date_start}至{date_end}）")
    set_run_font(run, "楷体_GB2312", "Times New Roman", 16, bold=False)

    # ===== 信息摘要 =====
    if summaries:
        # 摘要标签
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
        set_paragraph_spacing(p)
        run = p.add_run("【信息摘要】")
        set_run_font(run, "仿宋_GB2312", "Times New Roman", 16)

        # 摘要条目
        for i, s in enumerate(summaries):
            text = s if s.endswith("；") else s + "；"
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
            set_paragraph_spacing(p)
            run = p.add_run(f"{i+1}.{text}")
            set_run_font(run, "仿宋_GB2312", "Times New Roman", 16)

    # ===== 章节 =====
    global_news_idx = 0
    for si, sec in enumerate(sections):
        sec_num = CN_NUMS[si] if si < len(CN_NUMS) else str(si + 1)
        sec_title = sec.get("title", "")

        # 章节标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
        set_paragraph_spacing(p)
        run = p.add_run(f"{sec_num}、{sec_title}")
        set_run_font(run, "黑体", "黑体", 16)

        # 新闻条目
        for news in sec.get("newsItems", []):
            global_news_idx += 1
            headline = news.get("headline", "")
            body = news.get("body", "")

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
            set_paragraph_spacing(p)

            # 标题（加粗）
            run = p.add_run(f"{global_news_idx}.{headline}")
            set_run_font(run, "仿宋_GB2312", "Times New Roman", 16, bold=True)

            # 正文（不加粗）
            if body:
                run = p.add_run(body)
                set_run_font(run, "仿宋_GB2312", "Times New Roman", 16, bold=False)

    # ===== 写入到 BytesIO =====
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_filename(report_data: dict) -> str:
    """生成文件名"""
    date_start = report_data.get("dateStart", "")
    date_end = report_data.get("dateEnd", "")
    return f"国际移民一周资讯_{date_start}至{date_end}.docx".replace(" ", "")
